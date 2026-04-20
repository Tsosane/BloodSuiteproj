from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "FULL_SYSTEM_DOCUMENTATION.md"
TARGET = ROOT / "docs" / "Blood_Suite_Full_System_Documentation.docx"


def set_cell_text(cell, text):
    cell.text = text.strip()


def apply_page_borders(section):
    sect_pr = section._sectPr
    borders = sect_pr.find(qn("w:pgBorders"))
    if borders is None:
        borders = OxmlElement("w:pgBorders")
        borders.set(qn("w:offsetFrom"), "page")
        sect_pr.append(borders)

    for edge in ("top", "left", "bottom", "right"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "8")
        element.set(qn("w:space"), "24")
        element.set(qn("w:color"), "D32F2F")


def parse_markdown_table(lines, start_index):
    rows = []
    index = start_index

    while index < len(lines) and lines[index].strip().startswith("|"):
        rows.append(lines[index].strip())
        index += 1

    if len(rows) < 2:
        return None, start_index

    parsed_rows = []
    for row in rows:
        cells = [cell.strip() for cell in row.strip("|").split("|")]
        parsed_rows.append(cells)

    if all(re.fullmatch(r"[:\- ]+", cell or "") for cell in parsed_rows[1]):
        return parsed_rows[:1] + parsed_rows[2:], index

    return parsed_rows, index


def add_table(document, rows):
    if not rows:
        return

    column_count = max(len(row) for row in rows)
    table = document.add_table(rows=1, cols=column_count)
    table.style = "Table Grid"

    for cell_index, value in enumerate(rows[0]):
        set_cell_text(table.rows[0].cells[cell_index], value)

    for row in rows[1:]:
        cells = table.add_row().cells
        for cell_index, value in enumerate(row):
            set_cell_text(cells[cell_index], value)


def add_code_block(document, code_lines):
    for code_line in code_lines:
        paragraph = document.add_paragraph(style="No Spacing")
        run = paragraph.add_run(code_line.rstrip("\n"))
        run.font.name = "Consolas"
        run.font.size = Pt(9)


def build_document():
    markdown = SOURCE.read_text(encoding="utf-8").splitlines()

    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    apply_page_borders(section)

    normal_style = document.styles["Normal"]
    normal_style.font.name = "Calibri"
    normal_style.font.size = Pt(11)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Blood Suite\nFull System Documentation")
    run.bold = True
    run.font.size = Pt(22)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run("Generated from the repository documentation source")
    subtitle_run.italic = True
    subtitle_run.font.size = Pt(11)

    document.add_paragraph("")

    index = 0
    in_code_block = False
    code_lines = []

    while index < len(markdown):
        line = markdown[index]
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code_block:
                add_code_block(document, code_lines)
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
            index += 1
            continue

        if in_code_block:
            code_lines.append(line)
            index += 1
            continue

        if not stripped:
            document.add_paragraph("")
            index += 1
            continue

        if stripped.startswith("|"):
            rows, next_index = parse_markdown_table(markdown, index)
            if rows:
                add_table(document, rows)
                index = next_index
                continue

        heading_match = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if heading_match:
            level = min(len(heading_match.group(1)), 4)
            document.add_heading(heading_match.group(2).strip(), level=level)
            index += 1
            continue

        if re.match(r"^\d+\.\s+", stripped):
            paragraph = document.add_paragraph(style="List Number")
            paragraph.add_run(re.sub(r"^\d+\.\s+", "", stripped))
            index += 1
            continue

        if stripped.startswith("- "):
            paragraph = document.add_paragraph(style="List Bullet")
            paragraph.add_run(stripped[2:].strip())
            index += 1
            continue

        document.add_paragraph(stripped)
        index += 1

    if code_lines:
        add_code_block(document, code_lines)

    document.save(TARGET)


if __name__ == "__main__":
    build_document()
    print(TARGET)
